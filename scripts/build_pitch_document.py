"""
Builds docs/SmartTradeAI_Product_Overview.docx — a pitch/product deck for
sharing with the trading-influencer community: feature walkthrough with live
screenshots, differentiators, monetization models, and an enterprise /
future-roadmap section. Not part of the running app; run manually:

    python scripts/build_pitch_document.py
"""
from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "docs" / "pitch_screenshots"
OUT = ROOT / "docs" / "SmartTradeAI_Product_Overview.docx"

# ── Brand palette (matches the app's dark UI accent) ─────────────────────────
ACCENT = RGBColor(0x63, 0x66, 0xF1)     # indigo accent
ACCENT_DARK = RGBColor(0x43, 0x38, 0xCA)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xEF, 0x44, 0x44)
DARK = RGBColor(0x11, 0x14, 0x1F)
GREY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Base style tuning ─────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for i, size in zip(range(1, 4), (26, 18, 14)):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = ACCENT_DARK if i > 1 else DARK
    h.paragraph_format.space_before = Pt(18 if i == 1 else 12)
    h.paragraph_format.space_after = Pt(8)


def set_cell_shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_page_break():
    doc.add_page_break()


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text="", bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def add_bullets(items, bold_lead=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            r1.bold = True
            if rest:
                p.add_run(rest)
        else:
            p.add_run(item)


def add_screenshot(filename, caption=None, width_in=6.4):
    path = SHOTS / filename
    if not path.exists():
        add_para(f"[screenshot missing: {filename}]", italic=True, color=GREY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = GREY


def add_feature_section(title, icon_label, tagline, bullets, screenshot, caption=None):
    add_heading(title, level=2)
    add_para(tagline, italic=True, color=ACCENT_DARK, size=11.5)
    add_bullets(bullets)
    add_screenshot(screenshot, caption or title)
    add_para("", space_after=4)


def add_toc_field():
    """Insert a real Word TOC field (updates via right-click > Update Field,
    or automatically if the user accepts the update prompt on open)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and choose 'Update Field' to generate the Table of Contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(fld_text)
    r_element.append(fld_end)


def enable_update_fields_on_open():
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def kpi_row(cells_text):
    """A compact one-row 'stat strip' table for section highlights."""
    table = doc.add_table(rows=1, cols=len(cells_text))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for cell, (label, value) in zip(table.rows[0].cells, cells_text):
        cell.text = ""
        set_cell_shading(cell, "1E2433")
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(value)
        r1.bold = True
        r1.font.size = Pt(15)
        r1.font.color.rgb = WHITE
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0xAA, 0xB4, 0xC8)
    doc.add_paragraph()


def simple_table(headers, rows, col_widths_in=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], "4338CA")
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


# ═════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(140)
r = cover.add_run("SmartTrade AI")
r.font.size = Pt(48)
r.bold = True
r.font.color.rgb = ACCENT

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tag.add_run("AI-Powered Multi-Market Trading Intelligence Platform")
tr.font.size = Pt(18)
tr.font.color.rgb = GREY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(30)
sr = sub.add_run(
    "Real-time signals · AI/ML predictions · Multi-timeframe confluence\n"
    "Crypto · Forex · Commodities · Indian Equities · Indices"
)
sr.font.size = Pt(13)
sr.font.color.rgb = DARK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(60)
mr = meta.add_run(
    f"Product & Partnership Overview\n{datetime.date.today().strftime('%B %Y')}\nConfidential — Prepared for Partner Review"
)
mr.font.size = Pt(11)
mr.font.color.rgb = GREY
mr.italic = True

add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════
add_heading("Table of Contents", level=1)
add_toc_field()
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════
add_heading("Executive Summary", level=1)
add_para(
    "SmartTrade AI is a full-stack, AI-driven trading intelligence platform covering five asset "
    "classes — Crypto, Forex, Commodities, Indian Equities, and Indices — built entirely on free, "
    "no-API-key data sources (Delta Exchange India for live crypto, Yahoo Finance for everything else). "
    "It combines rule-based technical analysis, a proprietary multi-timeframe EMA confluence engine, "
    "and an ensemble of machine-learning models to generate, rank, and track trading signals in real time."
)
add_para(
    "The platform is already a working, end-to-end product — not a prototype. It includes live signal "
    "generation and auto-generation scheduling, a market scanner, backtesting engine, portfolio and "
    "watchlist tracking, a trade journal, risk-management tools, a full admin/audit layer, and a live "
    "broker connection (Delta Exchange India) for manual order execution. Every page shown in this "
    "document is a real, running screenshot of the live application — nothing here is a mockup."
)
kpi_row([
    ("Asset Classes Covered", "5"),
    ("Live Pages / Modules", "28+"),
    ("Signal Timeframes", "7"),
    ("Data Cost", "$0 / mo"),
])
add_para(
    "This document is intended for trading-community leaders, influencers, and potential partners "
    "evaluating SmartTrade AI for co-branded rollout, licensing, or enterprise/SaaS deployment to a "
    "trader audience. It walks through every feature with live screenshots, then lays out monetization "
    "paths and the roadmap for scaling this into a multi-tenant enterprise product."
)
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# PRODUCT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════
add_heading("Product Overview", level=1)
add_para(
    "SmartTrade AI is organized around a simple idea: turn raw market data into a ranked, explainable, "
    "risk-aware trade idea — automatically, continuously, across every market a trader cares about."
)
add_heading("How a signal is made", level=2)
add_bullets([
    ("Data ingestion — ", "live WebSocket prices from Delta Exchange India for crypto; Yahoo Finance for forex, commodities, Indian equities, and indices. No paid API keys required to run the platform."),
    ("Technical analysis engine — ", "a 7-stage pipeline (session gate, volatility regime, multi-timeframe trend gate, momentum gate, volume gate, confidence scoring, and structure-aware stop/target packaging) evaluates every asset on every timeframe from 1 minute to daily."),
    ("Proprietary EMA 9/21 multi-timeframe confluence — ", "each timeframe's EMA9/EMA21 cross is confirmed (or rejected) by the next-higher timeframe's own EMA9/21 read, with a built-in historical scrubber so the exact same logic can be back-tested bar-by-bar."),
    ("AI/ML ensemble — ", "Random Forest, XGBoost, LightGBM, and LSTM models vote on direction and confidence; predictions are cached and retrained automatically as new data arrives."),
    ("Risk-aware packaging — ", "every signal ships with an entry, stop-loss, and two take-profit levels, a computed risk:reward ratio, and a live confidence score — never a bare buy/sell flag."),
    ("Continuous operation — ", "a background scheduler regenerates signals on your configured schedule, tracks every open signal to its outcome (win/loss/expired), and feeds that history back into the win-rate and calibration analytics shown throughout the app."),
])
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# FEATURE WALKTHROUGH
# ═════════════════════════════════════════════════════════════════════════
add_heading("Feature Walkthrough", level=1)
add_para(
    "Every screenshot below is captured from the live, running application — populated with real "
    "positions, real signals, and real market data pulled at the time of writing.",
    italic=True, color=GREY,
)

add_feature_section(
    "Dashboard", "grid", "Command center: everything a trader needs on login, in one screen.",
    [
        "8 real-time KPI cards — expectancy, profit factor, active signals, open P&L, max drawdown, Sharpe ratio, win rate, and average R:R.",
        "AI Opportunity Radar — top 5 ranked ideas across all markets with live sparklines and confidence.",
        "Live Signals table with full trade plan (entry / SL / TP1 / TP2 / confidence / R:R / status) and per-signal AI model-agreement bars.",
        "AI Decision Inspector — click any signal to see exactly why the AI chose it (trend, momentum, volume, pattern, model votes) plus any risk warnings.",
        "Equity curve, win-rate-by-market, and confidence-calibration charts — the platform's own track record, transparently displayed.",
        "Full market heatmap with 5 view modes (price change, AI score, confidence, volatility, signal strength).",
    ],
    "dashboard.png",
)

add_feature_section(
    "Morning Briefing", "sunrise", "A one-screen pre-market read: regime, breadth, headlines, and key levels.",
    [
        "5 market-state cards — regime, volatility, sentiment, risk environment, global clarity.",
        "Market breadth donut (advancing / declining / unchanged) and a live sentiment gauge with by-market breakdown.",
        "Curated key headlines with sentiment tagging, and the day's economic calendar with impact levels.",
        "Overnight top gainers/losers and a crypto snapshot with 7-day trend sparklines.",
        "Configurable 'Key Levels to Watch' — pick any assets to track pivot levels (R2/R1/Pivot/S1/S2) with a live candlestick chart; the selection is saved per user.",
        "AI-generated insight bullets summarizing the session in plain English.",
    ],
    "morning_briefing.png",
)

add_feature_section(
    "Markets Overview (Crypto · Forex · Commodities · Indian Stocks · Indices)",
    "globe", "One unified market command center, five markets deep — switch instantly via tabs.",
    [
        "7 KPI cards scoped to the selected market: signals generated, active signals, win rate, avg R:R, open P&L, market sentiment, and volatility index.",
        "Market Sentiment gauge and breadth counts recalculate per-market the instant you switch tabs — never a stale 'all markets' number bleeding into a single-market view.",
        "AI Score heatmap covering every asset in the market (not just ones with an active signal), each with a Strong-Buy → Strong-Sell score.",
        "Top Opportunities, AI Performance (7-day win-rate/R:R chart), AI Model Consensus donut, News Impact, Upcoming Events, and one-click Quick Actions (Generate All, Run Scan, Manage Watchlist).",
    ],
    "markets_overview.png", "Markets Overview — Crypto tab (Forex / Commodities / Indian Stocks / Indices identical layout)",
)
add_screenshot("markets_forex.png", "Markets Overview — Forex")
add_screenshot("markets_commodities.png", "Markets Overview — Commodities")
add_screenshot("markets_indian_stocks.png", "Markets Overview — Indian Stocks")
add_screenshot("markets_indices.png", "Markets Overview — Indices")

add_feature_section(
    "All Signals — Live & History", "lightning", "Every signal the engine has ever produced, filterable and exportable.",
    [
        "Live tab: 14-column table (asset, market, signal, timeframe, entry, current, SL, TP1, TP2, confidence, R:R, age, regime, model-agreement bars, status).",
        "Open Positions live P&L panel — real-time distance to stop-loss and take-profit on every open signal.",
        "History tab: every closed signal with entry/exit/P&L%/outcome/confidence/duration — the full audit trail behind the win-rate numbers.",
        "Market / timeframe / signal-type / minimum-confidence filters, plus one-click CSV export.",
    ],
    "all_signals.png",
)
add_screenshot("signal_history.png", "All Signals — Signal History tab")

add_feature_section(
    "Auto Generate", "robot", "Set it and forget it: scheduled, hands-free signal generation.",
    [
        "Configure markets, timeframes (multi-select), signal filter, minimum confidence, and an optional asset whitelist.",
        "Auto-repeat on any interval from 1 minute to 4 hours, with Start / Pause / Stop controls and a live run log.",
        "Session stats (generated / runs / errors / buy / sell / hold) and a live countdown to the next run.",
        "Market Sentiment, AI Prediction shortcut, and Quick Stats panel alongside the control console.",
    ],
    "auto_generate.png",
)

add_feature_section(
    "Market Scanner", "search", "Screen every market against technical filters in real time.",
    [
        "13 one-click filters — Strong Buy/Sell, Breakout/Breakdown, Volume Spike, RSI Oversold/Overbought, Gap Up/Down, 52-Week High/Low.",
        "Market / timeframe / minimum-confidence controls, with results ranked by a derived confidence score.",
        "Per-row matched-filter icons, one-click watchlist add, and CSV export of scan results.",
        "7 KPI cards summarizing the scan: results, strong signals, win rate, avg R:R, open P&L, sentiment, and volatility.",
    ],
    "scanner.png",
)

add_feature_section(
    "Market Heatmap", "grid-3x3", "The whole market's temperature, at a glance.",
    [
        "Grouped by market with color-coded tiles (price change, with a legend from Strong Buy to Strong Sell).",
        "One click through to full asset detail from any tile.",
    ],
    "market_heatmap.png",
)

add_feature_section(
    "Backtesting", "clock-history", "Validate a strategy against history before risking a rupee.",
    [
        "Configure asset, strategy parameters, timeframe, and date range; run and get equity curve, win rate, total trades, profit factor, and max drawdown.",
        "Full trade-by-trade ledger for auditability.",
    ],
    "backtesting.png",
)

add_feature_section(
    "AI Insights", "cpu", "Ensemble machine-learning predictions, explained — not a black box.",
    [
        "Per-asset, per-timeframe bullish/bearish probability and confidence from the ML ensemble (Random Forest + XGBoost + LightGBM + LSTM).",
        "Confluence summary across multiple selected timeframes — agreement across timeframes is flagged as a high-conviction setup.",
    ],
    "ai_insights.png",
)

add_feature_section(
    "TA Summary — including the proprietary EMA 9/21 Multi-Timeframe engine", "speedometer2",
    "Four ways to read the technical picture: rule-based ratings, AI ratings, live quotes, and our own EMA9/21 MTF confluence method.",
    [
        "Technical Ratings — 12-indicator consensus rating per asset, per timeframe (5m through daily) plus an overall score.",
        "AI Ratings — the same grid powered by the ML ensemble instead of rule-based indicators.",
        "Live Quotes — real-time OHLCV for every tracked asset.",
        "EMA 9/21 MTF (new, proprietary) — reads EMA9 vs EMA21 on each timeframe, confirmed by price reclaiming the fast EMA, then cross-checked against the same read on the next-higher timeframe (5m→15m→30m→1h→2h→4h→1d). Agreement = Strong Buy/Sell; disagreement is honestly reported as Neutral rather than guessing.",
        "Built-in history scrubber — step back through past bars (1/5/10/20/50 at a time, or jump to an exact offset) to see exactly what the EMA9/21 MTF grid showed at any past moment, for pattern discovery and manual backtesting. Every read is point-in-time correct — it is architecturally impossible for a historical read to see a future bar.",
        "Click any cell for a full breakdown: the exact EMA9, EMA21, and price on both timeframes and the plain-English reasoning behind the rating.",
    ],
    "ta_summary.png",
)

add_feature_section(
    "MTF Analysis", "table", "Multi-timeframe trend alignment in one grid.",
    ["Every asset × every timeframe, color-coded by trend agreement — green rows are high-confidence aligned setups, red rows are aligned bearish, mixed rows flag indecision."],
    "mtf_analysis.png",
)

add_feature_section(
    "Model Performance", "bar-chart", "The AI's own report card — nothing hidden.",
    ["Accuracy/precision/recall per model, performance-over-time trend, and a per-market breakdown — full transparency on how the ML ensemble is actually performing, not just what it predicts next."],
    "model_performance.png",
)

add_feature_section(
    "Advanced Analysis", "layers", "Professional-grade charting for the assets that matter most.",
    ["Large candlestick chart with support/resistance and Fibonacci overlays, a liquidity-pool table, and timeframe tabs for deep single-asset study."],
    "advanced_analysis.png",
)

add_feature_section(
    "Analytics", "graph-up", "Roll up performance by market, signal type, and timeframe.",
    ["Win-rate-by-market bar chart, signal-type distribution donut, and a performance-by-timeframe table — the numbers behind every claim the platform makes."],
    "analytics.png",
)

add_feature_section(
    "Market News", "newspaper", "Curated, sentiment-tagged news — no tab-switching to a broker terminal.",
    ["Card grid of headlines with source, timestamp, AI-derived sentiment/impact tags, and a direct read-more link."],
    "news.png",
)

add_feature_section(
    "Economic Calendar", "calendar-event", "Never get caught by a surprise NFP or FOMC print again.",
    ["Full schedule of macro events with country, impact level (High/Medium/Low), and forecast/previous/actual values, correctly timezone-converted to IST."],
    "economic_calendar.png",
)

add_feature_section(
    "Portfolio", "pie-chart", "Track real holdings and live, unrealized P&L in one place.",
    ["KPI cards for total invested, current value, total P&L, and P&L%; a full holdings table with live prices, per-position P&L, and days held; Add Position and CSV export."],
    "portfolio.png",
)

add_feature_section(
    "Watchlist", "star", "Your shortlist, with live prices and one-click alerts.",
    ["Saved assets with live price, % change, and a confluence/signal badge; alert-price triggers fire a notification the moment they're crossed."],
    "watchlist.png",
)

add_feature_section(
    "Trading (Live Broker Connection)", "wallet2", "Bring your own Delta Exchange India account — fully non-custodial.",
    [
        "Direct integration with Delta Exchange India for manual order placement, balances, positions, and order history.",
        "Non-custodial by design — the platform never holds funds; each user connects their own broker API key with trading permission, giving full control and a clean compliance story for an enterprise rollout.",
        "A clear, guided connection flow (shown below) when no broker is yet linked — no confusing dead ends.",
    ],
    "trading.png",
)

add_feature_section(
    "Risk Manager", "shield-check", "Position sizing and risk:reward, calculated instantly.",
    [
        "Account size and risk-per-trade inputs feed a live risk-context KPI strip (account size, risk %, max risk per trade, historical win rate).",
        "Trade calculator: entry, stop-loss, target → instant position size, position value, and risk:reward, with quick-load from any active signal.",
    ],
    "risk_manager.png",
)

add_feature_section(
    "My Performance", "graph-up-arrow", "The trader's own scoreboard.",
    ["Equity curve, win rate, average win/loss, profit factor, and best/worst trade — a personal track record separate from the platform-wide analytics."],
    "my_performance.png",
)

add_feature_section(
    "Trade Journal", "journal-text", "Discipline, systematized.",
    [
        "Full trade log with date, asset, direction, timeframe, entry/exit, P&L, outcome, and emotion tagging.",
        "Trading Insights panel and a Weekly Review (day-by-day trades/P&L/win-loss plus a notes field) to build a real self-review habit.",
    ],
    "trade_journal.png",
)

add_feature_section(
    "Settings", "gear", "Full account control in one place.",
    ["Profile, notification preferences (email/Telegram), password/2FA, and per-user asset selection for which symbols appear across TA Summary, MTF Analysis, and more."],
    "settings.png",
)

add_feature_section(
    "Help & FAQ", "question-circle", "A living methodology reference — not just a support page.",
    ["Jump-to-section navigation across every feature (Signals, AI Insights, TA Summary, MTF, Scanner, Backtesting, Risk Manager, Trade Journal, Briefing, Portfolio, Watchlist, Settings) with exact methodology explanations for each."],
    "help_faq.png",
)

add_feature_section(
    "Asset Detail", "graph-up", "The single-asset deep-dive every trader lands on.",
    [
        "Live candlestick chart with timeframe switcher, key indicators panel, and tabs for Chart / Indicators / Signals / Signals History / Profile.",
        "Sidebar: market sentiment gauge, live price, AI prediction gauge, and quick stats (exchange, data source, status).",
        "Bottom row: market heatmap, top opportunities, and 7-day AI performance — full market context without leaving the asset page.",
    ],
    "asset_detail.png",
)

add_page_break()
add_heading("Admin & Enterprise Control Panel", level=2)
add_para(
    "SmartTrade AI ships with a genuine admin layer — not an afterthought. This is the foundation the "
    "enterprise/multi-tenant roadmap (Section 7) builds on."
)
add_screenshot("admin_dashboard.png", "Admin Dashboard — system health, user management, platform analytics")
add_bullets([
    "Live system health: CPU / memory / disk usage and DB status.",
    "User management — role, status, join date, enable/disable, per-user activity.",
    "API Configs — manage every connected data/exchange provider, with live status per feed.",
    "Asset management and full, exportable audit logs (every login, config change, and admin action).",
])
add_screenshot("admin_users.png", "Admin — Users")
add_screenshot("admin_api_configs.png", "Admin — API Configs")
add_screenshot("admin_assets.png", "Admin — Assets")
add_screenshot("admin_logs.png", "Admin — System Logs")

add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# DIFFERENTIATORS
# ═════════════════════════════════════════════════════════════════════════
add_heading("Why This Wins", level=1)
add_bullets([
    ("Zero data-cost architecture — ", "every price feed is free (Delta Exchange India + Yahoo Finance). No Bloomberg terminal, no paid API tier, no per-seat data licensing eating into margin."),
    ("A genuinely proprietary signal method — ", "the EMA 9/21 multi-timeframe confluence engine, with point-in-time-correct historical scrubbing, is not a repackaged open-source indicator — it was purpose-built for this platform and is explainable cell-by-cell."),
    ("Explainability everywhere — ", "the AI Decision Inspector, per-cell EMA breakdowns, and model-performance transparency mean this never feels like a black box — a major trust unlock for a skeptical trader audience."),
    ("Full lifecycle, not just signals — ", "portfolio, watchlist, risk sizing, trade journal, and performance analytics mean a user never has to leave the platform to manage the trade the signal started."),
    ("Non-custodial live execution — ", "the Delta Exchange India connection is bring-your-own-key; the platform never touches user funds, which massively simplifies compliance for a commercial rollout."),
    ("Enterprise-ready from day one — ", "role-based access, a full audit log, and a Docker-based deployment path already exist — this is not a green-field engineering effort to make it enterprise-sellable."),
])
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# MONETIZATION MODELS
# ═════════════════════════════════════════════════════════════════════════
add_heading("Monetization Models", level=1)
add_para(
    "The data model already includes Free / Premium / Admin subscription tiers with configurable signal "
    "delay, watchlist limits, and feature gates — the plumbing for a paid tier already exists."
)
simple_table(
    ["Model", "How it works", "Best fit"],
    [
        ("Direct subscription (B2C)", "Free tier (delayed signals, capped watchlist) → Premium (real-time signals, AI Insights, backtesting, unlimited watchlist).", "Selling directly to a trader audience via an influencer's community"),
        ("Influencer co-branded / white-label", "Re-skin with the influencer's brand; they promote it to their audience for a revenue share or flat licensing fee.", "Trading YouTubers, Telegram/Discord community leads"),
        ("Signal / data API licensing", "License the signal-generation engine's output (or the EMA9/21 MTF confluence feed specifically) to other apps or bots via API.", "Other fintech builders, prop-trading tool vendors"),
        ("Enterprise / prop-firm SaaS", "Multi-seat deployment for a trading firm or prop-trading community, with per-seat admin controls and audit logging already built in.", "Prop trading firms, trading academies"),
        ("Affiliate / broker partnership", "Delta Exchange India integration is a natural affiliate-revenue hook — referral commission on connected accounts.", "Broker partnerships"),
        ("Paid add-on modules", "Backtesting, AI Insights, and Advanced Analysis are already gated as 'PRO' — sell as an upsell inside any tier.", "Upsell within any of the above"),
    ],
    col_widths_in=[1.7, 3.3, 1.5],
)
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# ENTERPRISE DEPLOYMENT ROADMAP
# ═════════════════════════════════════════════════════════════════════════
add_heading("Enterprise Deployment Roadmap", level=1)
add_para(
    "SmartTrade AI runs today as a single-tenant Flask application with SQLite, a background scheduler, "
    "and a WebSocket price stream — already Dockerized. Scaling this to serve many organizations or a large "
    "paid user base is a staged infrastructure evolution, not a rewrite."
)
add_heading("Phase 1 — Multi-user hardening (near-term)", level=2)
add_bullets([
    "Migrate from SQLite to PostgreSQL for concurrent-write safety at scale.",
    "Move the in-memory job scheduler to a durable task queue (Celery/RQ) so signal generation survives restarts and scales horizontally.",
    "Introduce Redis for caching and rate-limit storage (already abstracted behind Flask-Caching/Flask-Limiter — a config change, not a rewrite).",
])
add_heading("Phase 2 — True multi-tenancy", level=2)
add_bullets([
    "Organization/workspace model on top of the existing Role/Subscription tables — each tenant gets isolated users, watchlists, portfolios, and API configs.",
    "Per-tenant branding (logo, color theme, domain) building on the theme system already in the frontend.",
    "Tenant-level admin dashboards, reusing the existing Admin Panel/Audit Log architecture.",
])
add_heading("Phase 3 — Enterprise operations", level=2)
add_bullets([
    "SSO / SAML / enterprise auth alongside the existing JWT + 2FA login.",
    "Horizontal scaling of the signal engine and market-data collectors behind a load balancer; containerized deployment via the existing Dockerfile/docker-compose, promoted to Kubernetes for larger tenants.",
    "SLA-grade monitoring, alerting, and on-call runbooks around the existing structured logging and audit trail.",
    "Formal data-retention and compliance policies (the audit log and non-custodial broker design already give this a strong starting point).",
])
add_heading("Phase 4 — Platform expansion", level=2)
add_bullets([
    "Additional broker integrations beyond Delta Exchange India (multi-broker abstraction layer).",
    "Public signal/data API product (rate-limited, API-key-gated) as a standalone revenue line.",
    "Mobile app (iOS/Android) consuming the same API layer.",
])
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# FUTURE ENHANCEMENTS
# ═════════════════════════════════════════════════════════════════════════
add_heading("Future Enhancements Roadmap", level=1)
add_heading("Near-term (0–3 months)", level=2)
add_bullets([
    "More proprietary confluence methods alongside EMA 9/21 MTF (e.g., volume-weighted MTF confirmation, structure-based break-of-structure detection).",
    "Telegram/WhatsApp push notifications for high-confidence signals (Telegram hook already exists in the backend).",
    "Expanded position-sizing methods in Risk Manager (ATR-based, volatility-targeted, fractional Kelly with caps).",
    "Deeper backtesting: walk-forward validation and Monte Carlo stress testing (reshuffled trade order, slippage shocks) to report a probability-of-ruin alongside the equity curve.",
])
add_heading("Mid-term (3–9 months)", level=2)
add_bullets([
    "Paper-trading mode that runs the exact same signal → risk → execution pipeline as live trading, with simulated fills, for safe strategy validation before going live.",
    "Model registry and drift detection for the AI ensemble — formal candidate → shadow → live promotion workflow with automatic rollback on performance decay.",
    "Portfolio-level risk engine: correlation/concentration limits across positions, not just per-trade sizing.",
    "Additional broker integrations (multi-broker order routing) and options/futures support alongside spot.",
])
add_heading("Long-term (9+ months)", level=2)
add_bullets([
    "AI agent layer — specialized agents (market scanner, regime classifier, risk reviewer, post-trade critic) that debate a high-confidence signal before it's surfaced, with a final aggregator recommendation — always subordinate to the deterministic risk engine, never bypassing it.",
    "Social/copy-trading layer — opt-in strategy sharing and leaderboard, a natural fit for an influencer-led community rollout.",
    "Native mobile apps and a public developer API for third-party bot builders.",
    "Prop-firm challenge tracker — rule-based evaluation tracking (drawdown limits, profit targets) for trading-challenge communities.",
])
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# TECHNOLOGY STACK
# ═════════════════════════════════════════════════════════════════════════
add_heading("Technology Stack", level=1)
simple_table(
    ["Layer", "Technology"],
    [
        ("Backend", "Python, Flask, SQLAlchemy, Flask-SocketIO, APScheduler, Flask-JWT-Extended"),
        ("AI / ML", "scikit-learn (Random Forest), XGBoost, LightGBM, LSTM ensemble"),
        ("Market data", "Delta Exchange India (live crypto WebSocket), Yahoo Finance (forex/commodities/equities/indices) — both free, no API key"),
        ("Frontend", "Server-rendered Jinja templates, Chart.js, vanilla JS, light/dark theme system"),
        ("Database", "SQLite (current), PostgreSQL-ready migration path"),
        ("Deployment", "Dockerfile + docker-compose already included"),
        ("Security", "JWT auth with 2FA, encrypted API-key storage, full audit logging, rate limiting"),
    ],
    col_widths_in=[1.8, 4.7],
)
add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# CLOSING
# ═════════════════════════════════════════════════════════════════════════
add_heading("Next Steps", level=1)
add_para(
    "SmartTrade AI is a working, feature-complete platform today, with a clear, staged path to enterprise "
    "multi-tenant deployment and several monetization models that don't require re-engineering the core "
    "product. We're looking to partner with trading-community leaders and influencers who want to bring a "
    "genuinely capable, explainable AI trading platform to their audience — whether that's white-label "
    "licensing, a revenue-share subscription rollout, or a straightforward data/signal API partnership."
)
add_para("Let's talk about the right model for your community.", bold=True, size=13, color=ACCENT_DARK)

enable_update_fields_on_open()
doc.save(OUT)
print(f"Saved: {OUT}")
